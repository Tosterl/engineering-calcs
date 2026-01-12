"""
Report generation service for Engineering Calculations Database.

This module provides comprehensive report generation capabilities for engineering
calculations, supporting PDF and Word document formats with professional formatting.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional, Union

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.core.calculations import CalculationResult
from src.core.units import Quantity


@dataclass
class ReportOptions:
    """
    Configuration options for report generation.

    Attributes:
        title: The report title.
        project_name: Name of the project.
        author: Author of the report.
        include_charts: Whether to include charts in the report.
        include_steps: Whether to include intermediate calculation steps.
        logo_path: Optional path to a logo image file.
        company_name: Optional company name for the header.
    """
    title: str
    project_name: str
    author: str
    include_charts: bool = True
    include_steps: bool = True
    logo_path: Optional[str] = None
    company_name: Optional[str] = None


class ReportService:
    """
    Service for generating professional engineering calculation reports.

    This service provides methods to generate PDF and Word documents from
    calculation results, with support for custom formatting options.
    """

    # Default styles
    HEADER_COLOR = colors.HexColor("#1a365d")
    ACCENT_COLOR = colors.HexColor("#2b6cb0")
    TABLE_HEADER_BG = colors.HexColor("#e2e8f0")
    TABLE_ALT_ROW_BG = colors.HexColor("#f7fafc")

    def __init__(self) -> None:
        """Initialize the report service with default settings."""
        self._pdf_styles = self._create_pdf_styles()

    def _create_pdf_styles(self) -> Dict[str, ParagraphStyle]:
        """Create custom PDF paragraph styles."""
        styles = getSampleStyleSheet()

        custom_styles = {
            "ReportTitle": ParagraphStyle(
                "ReportTitle",
                parent=styles["Title"],
                fontSize=24,
                textColor=self.HEADER_COLOR,
                spaceAfter=20,
                alignment=1,  # Center
            ),
            "SectionHeader": ParagraphStyle(
                "SectionHeader",
                parent=styles["Heading1"],
                fontSize=14,
                textColor=self.HEADER_COLOR,
                spaceBefore=20,
                spaceAfter=10,
                borderWidth=1,
                borderColor=self.ACCENT_COLOR,
                borderPadding=5,
            ),
            "SubsectionHeader": ParagraphStyle(
                "SubsectionHeader",
                parent=styles["Heading2"],
                fontSize=12,
                textColor=self.ACCENT_COLOR,
                spaceBefore=15,
                spaceAfter=8,
            ),
            "BodyText": ParagraphStyle(
                "BodyText",
                parent=styles["Normal"],
                fontSize=10,
                spaceBefore=6,
                spaceAfter=6,
            ),
            "Formula": ParagraphStyle(
                "Formula",
                parent=styles["Normal"],
                fontSize=10,
                fontName="Courier",
                leftIndent=20,
                spaceBefore=4,
                spaceAfter=4,
                backColor=colors.HexColor("#f0f4f8"),
            ),
            "MetaInfo": ParagraphStyle(
                "MetaInfo",
                parent=styles["Normal"],
                fontSize=9,
                textColor=colors.gray,
            ),
        }

        return custom_styles

    @staticmethod
    def _format_value(value: Any, unit: Optional[str] = None) -> str:
        """
        Format a value with its unit for display.

        Args:
            value: The value to format (can be Quantity, float, int, or str).
            unit: Optional unit string if value is not a Quantity.

        Returns:
            Formatted string representation of the value.
        """
        if isinstance(value, Quantity):
            return value.format()
        elif isinstance(value, float):
            # Format floats with reasonable precision
            if abs(value) >= 1000 or (abs(value) < 0.01 and value != 0):
                formatted = f"{value:.4e}"
            else:
                formatted = f"{value:.4f}".rstrip("0").rstrip(".")
            if unit:
                return f"{formatted} {unit}"
            return formatted
        elif isinstance(value, int):
            if unit:
                return f"{value} {unit}"
            return str(value)
        else:
            return str(value)

    def _create_header(
        self,
        elements: List[Any],
        title: str,
        date: datetime,
        options: Optional[ReportOptions] = None,
    ) -> None:
        """
        Create the report header section for PDF.

        Args:
            elements: List to append header elements to.
            title: The report title.
            date: The report date.
            options: Optional report configuration options.
        """
        # Logo placeholder (if provided)
        if options and options.logo_path:
            try:
                logo = Image(options.logo_path, width=1.5 * inch, height=0.75 * inch)
                elements.append(logo)
                elements.append(Spacer(1, 10))
            except Exception:
                # If logo cannot be loaded, skip it
                pass

        # Company name (if provided)
        if options and options.company_name:
            company_para = Paragraph(
                options.company_name,
                self._pdf_styles["MetaInfo"],
            )
            elements.append(company_para)
            elements.append(Spacer(1, 5))

        # Title
        title_para = Paragraph(title, self._pdf_styles["ReportTitle"])
        elements.append(title_para)

        # Project name and metadata
        meta_parts = []
        if options and options.project_name:
            meta_parts.append(f"Project: {options.project_name}")
        if options and options.author:
            meta_parts.append(f"Author: {options.author}")
        meta_parts.append(f"Date: {date.strftime('%Y-%m-%d %H:%M')}")

        for meta in meta_parts:
            meta_para = Paragraph(meta, self._pdf_styles["MetaInfo"])
            elements.append(meta_para)

        elements.append(Spacer(1, 20))

        # Horizontal line
        line_table = Table([[""]], colWidths=[6.5 * inch])
        line_table.setStyle(TableStyle([
            ("LINEABOVE", (0, 0), (-1, 0), 2, self.ACCENT_COLOR),
        ]))
        elements.append(line_table)
        elements.append(Spacer(1, 20))

    def _add_table(
        self,
        elements: List[Any],
        headers: List[str],
        rows: List[List[str]],
        col_widths: Optional[List[float]] = None,
    ) -> None:
        """
        Add a formatted table to the PDF elements.

        Args:
            elements: List to append the table to.
            headers: List of column header strings.
            rows: List of row data (each row is a list of strings).
            col_widths: Optional list of column widths in inches.
        """
        if not rows:
            return

        # Combine headers and rows
        table_data = [headers] + rows

        # Calculate column widths if not provided
        if col_widths is None:
            num_cols = len(headers)
            col_widths = [6.5 * inch / num_cols] * num_cols

        table = Table(table_data, colWidths=col_widths)

        # Style the table
        style = TableStyle([
            # Header styling
            ("BACKGROUND", (0, 0), (-1, 0), self.TABLE_HEADER_BG),
            ("TEXTCOLOR", (0, 0), (-1, 0), self.HEADER_COLOR),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 10),
            ("TOPPADDING", (0, 0), (-1, 0), 10),
            # Body styling
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 9),
            ("ALIGN", (0, 1), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
            ("TOPPADDING", (0, 1), (-1, -1), 6),
            # Grid
            ("GRID", (0, 0), (-1, -1), 0.5, colors.gray),
            ("BOX", (0, 0), (-1, -1), 1, self.ACCENT_COLOR),
        ])

        # Alternate row colors
        for i in range(1, len(table_data)):
            if i % 2 == 0:
                style.add("BACKGROUND", (0, i), (-1, i), self.TABLE_ALT_ROW_BG)

        table.setStyle(style)
        elements.append(table)
        elements.append(Spacer(1, 15))

    def generate_pdf(
        self,
        calculation_result: CalculationResult,
        output_path: str,
        options: Optional[Union[ReportOptions, dict]] = None,
    ) -> str:
        """
        Generate a PDF report from a calculation result.

        Args:
            calculation_result: The calculation result to document.
            output_path: Path where the PDF file should be saved.
            options: Report configuration options (ReportOptions or dict).

        Returns:
            The path to the generated PDF file.

        Raises:
            IOError: If the file cannot be written.
        """
        # Convert dict options to ReportOptions if needed
        if isinstance(options, dict):
            options = ReportOptions(
                title=options.get("title", "Engineering Calculation Report"),
                project_name=options.get("project_name", ""),
                author=options.get("author", ""),
                include_charts=options.get("include_charts", True),
                include_steps=options.get("include_steps", True),
                logo_path=options.get("logo_path"),
                company_name=options.get("company_name"),
            )
        elif options is None:
            options = ReportOptions(
                title="Engineering Calculation Report",
                project_name="",
                author="",
            )

        # Create the PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )

        elements: List[Any] = []

        # Header
        self._create_header(
            elements,
            options.title,
            calculation_result.timestamp,
            options,
        )

        # Calculation name and description
        calc_name = calculation_result.calculation_name or "Calculation"
        description = calculation_result.metadata.get("description", "")

        elements.append(
            Paragraph(f"Calculation: {calc_name}", self._pdf_styles["SectionHeader"])
        )
        if description:
            elements.append(
                Paragraph(description, self._pdf_styles["BodyText"])
            )
        elements.append(Spacer(1, 10))

        # Input parameters table
        elements.append(
            Paragraph("Input Parameters", self._pdf_styles["SubsectionHeader"])
        )
        input_headers = ["Parameter", "Value", "Unit"]
        input_rows = []
        for name, value in calculation_result.inputs.items():
            if isinstance(value, Quantity):
                input_rows.append([name, f"{value.magnitude:.4g}", value.unit_string])
            else:
                input_rows.append([name, self._format_value(value), "-"])

        self._add_table(
            elements,
            input_headers,
            input_rows,
            col_widths=[2.5 * inch, 2 * inch, 2 * inch],
        )

        # Calculation steps (if enabled)
        if options.include_steps and calculation_result.intermediate_steps:
            elements.append(
                Paragraph("Calculation Steps", self._pdf_styles["SubsectionHeader"])
            )

            for i, step in enumerate(calculation_result.intermediate_steps, 1):
                step_text = f"<b>Step {i}:</b> {step.description}"
                elements.append(Paragraph(step_text, self._pdf_styles["BodyText"]))

                if step.formula:
                    elements.append(
                        Paragraph(f"Formula: {step.formula}", self._pdf_styles["Formula"])
                    )

                if step.substitution:
                    elements.append(
                        Paragraph(
                            f"Substitution: {step.substitution}",
                            self._pdf_styles["Formula"],
                        )
                    )

                result_str = self._format_value(step.result)
                elements.append(
                    Paragraph(f"Result: {result_str}", self._pdf_styles["Formula"])
                )
                elements.append(Spacer(1, 8))

        # Results table
        elements.append(
            Paragraph("Results", self._pdf_styles["SubsectionHeader"])
        )
        result_headers = ["Output", "Value", "Unit"]
        result_rows = []
        for name, value in calculation_result.outputs.items():
            if isinstance(value, Quantity):
                result_rows.append([name, f"{value.magnitude:.4g}", value.unit_string])
            else:
                result_rows.append([name, self._format_value(value), "-"])

        self._add_table(
            elements,
            result_headers,
            result_rows,
            col_widths=[2.5 * inch, 2 * inch, 2 * inch],
        )

        # Charts placeholder (if enabled and provided)
        if options.include_charts:
            charts = calculation_result.metadata.get("charts", [])
            if charts:
                elements.append(
                    Paragraph("Charts", self._pdf_styles["SubsectionHeader"])
                )
                for chart_path in charts:
                    try:
                        chart_img = Image(chart_path, width=5 * inch, height=3 * inch)
                        elements.append(chart_img)
                        elements.append(Spacer(1, 10))
                    except Exception:
                        # Skip charts that cannot be loaded
                        pass

        # References
        references = calculation_result.metadata.get("references", [])
        if references:
            elements.append(
                Paragraph("References", self._pdf_styles["SubsectionHeader"])
            )
            for i, ref in enumerate(references, 1):
                ref_text = f"{i}. {ref}"
                elements.append(Paragraph(ref_text, self._pdf_styles["BodyText"]))

        # Build the PDF
        doc.build(elements)

        return output_path

    def generate_word(
        self,
        calculation_result: CalculationResult,
        output_path: str,
        options: Optional[Union[ReportOptions, dict]] = None,
    ) -> str:
        """
        Generate a Word document report from a calculation result.

        Args:
            calculation_result: The calculation result to document.
            output_path: Path where the Word file should be saved.
            options: Report configuration options (ReportOptions or dict).

        Returns:
            The path to the generated Word file.

        Raises:
            IOError: If the file cannot be written.
        """
        # Convert dict options to ReportOptions if needed
        if isinstance(options, dict):
            options = ReportOptions(
                title=options.get("title", "Engineering Calculation Report"),
                project_name=options.get("project_name", ""),
                author=options.get("author", ""),
                include_charts=options.get("include_charts", True),
                include_steps=options.get("include_steps", True),
                logo_path=options.get("logo_path"),
                company_name=options.get("company_name"),
            )
        elif options is None:
            options = ReportOptions(
                title="Engineering Calculation Report",
                project_name="",
                author="",
            )

        # Create the Word document
        doc = Document()

        # Set up custom styles
        self._setup_word_styles(doc)

        # Logo placeholder
        if options.logo_path:
            try:
                doc.add_picture(options.logo_path, width=Inches(1.5))
            except Exception:
                pass

        # Company name
        if options.company_name:
            company_para = doc.add_paragraph(options.company_name)
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            company_para.style = doc.styles["Caption"]

        # Title
        title_para = doc.add_heading(options.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if options.project_name:
            meta_para.add_run(f"Project: {options.project_name}\n")
        if options.author:
            meta_para.add_run(f"Author: {options.author}\n")
        meta_para.add_run(
            f"Date: {calculation_result.timestamp.strftime('%Y-%m-%d %H:%M')}"
        )

        doc.add_paragraph()  # Spacer

        # Calculation name and description
        calc_name = calculation_result.calculation_name or "Calculation"
        description = calculation_result.metadata.get("description", "")

        doc.add_heading(f"Calculation: {calc_name}", level=1)
        if description:
            doc.add_paragraph(description)

        # Input parameters table
        doc.add_heading("Input Parameters", level=2)
        input_headers = ["Parameter", "Value", "Unit"]
        input_rows = []
        for name, value in calculation_result.inputs.items():
            if isinstance(value, Quantity):
                input_rows.append([name, f"{value.magnitude:.4g}", value.unit_string])
            else:
                input_rows.append([name, self._format_value(value), "-"])

        self._add_word_table(doc, input_headers, input_rows)

        # Calculation steps
        if options.include_steps and calculation_result.intermediate_steps:
            doc.add_heading("Calculation Steps", level=2)

            for i, step in enumerate(calculation_result.intermediate_steps, 1):
                step_para = doc.add_paragraph()
                step_para.add_run(f"Step {i}: ").bold = True
                step_para.add_run(step.description)

                if step.formula:
                    formula_para = doc.add_paragraph(f"Formula: {step.formula}")
                    formula_para.style = doc.styles["Quote"]

                if step.substitution:
                    sub_para = doc.add_paragraph(f"Substitution: {step.substitution}")
                    sub_para.style = doc.styles["Quote"]

                result_str = self._format_value(step.result)
                result_para = doc.add_paragraph(f"Result: {result_str}")
                result_para.style = doc.styles["Quote"]

                doc.add_paragraph()  # Spacer

        # Results table
        doc.add_heading("Results", level=2)
        result_headers = ["Output", "Value", "Unit"]
        result_rows = []
        for name, value in calculation_result.outputs.items():
            if isinstance(value, Quantity):
                result_rows.append([name, f"{value.magnitude:.4g}", value.unit_string])
            else:
                result_rows.append([name, self._format_value(value), "-"])

        self._add_word_table(doc, result_headers, result_rows)

        # Charts
        if options.include_charts:
            charts = calculation_result.metadata.get("charts", [])
            if charts:
                doc.add_heading("Charts", level=2)
                for chart_path in charts:
                    try:
                        doc.add_picture(chart_path, width=Inches(5))
                        doc.add_paragraph()  # Spacer
                    except Exception:
                        pass

        # References
        references = calculation_result.metadata.get("references", [])
        if references:
            doc.add_heading("References", level=2)
            for i, ref in enumerate(references, 1):
                doc.add_paragraph(f"{i}. {ref}")

        # Save the document
        doc.save(output_path)

        return output_path

    def _setup_word_styles(self, doc: Document) -> None:
        """
        Set up custom styles for the Word document.

        Args:
            doc: The Word document to configure styles for.
        """
        styles = doc.styles

        # Modify Title style
        title_style = styles["Title"]
        title_style.font.size = Pt(24)
        title_style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        title_style.font.bold = True

        # Modify Heading 1 style
        heading1_style = styles["Heading 1"]
        heading1_style.font.size = Pt(14)
        heading1_style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

        # Modify Heading 2 style
        heading2_style = styles["Heading 2"]
        heading2_style.font.size = Pt(12)
        heading2_style.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

    def _add_word_table(
        self,
        doc: Document,
        headers: List[str],
        rows: List[List[str]],
    ) -> None:
        """
        Add a formatted table to the Word document.

        Args:
            doc: The Word document to add the table to.
            headers: List of column header strings.
            rows: List of row data.
        """
        if not rows:
            return

        # Create table with header row and data rows
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Style header cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            # Set background color (requires python-docx-ng or additional code)
            # Note: Basic python-docx doesn't support cell shading easily

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                row.cells[col_idx].text = cell_text

        doc.add_paragraph()  # Spacer after table

    def generate_summary(self, results: List[CalculationResult]) -> str:
        """
        Generate a text summary of multiple calculation results.

        Args:
            results: List of calculation results to summarize.

        Returns:
            A formatted text summary string.
        """
        if not results:
            return "No calculations to summarize."

        summary_parts = [
            "=" * 60,
            "ENGINEERING CALCULATIONS SUMMARY",
            "=" * 60,
            f"Total Calculations: {len(results)}",
            f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
            "-" * 60,
            "",
        ]

        for i, result in enumerate(results, 1):
            calc_name = result.calculation_name or f"Calculation {i}"
            category = result.metadata.get("category", "Uncategorized")
            description = result.metadata.get("description", "")

            summary_parts.append(f"{i}. {calc_name}")
            summary_parts.append(f"   Category: {category}")
            if description:
                summary_parts.append(f"   Description: {description}")
            summary_parts.append(f"   Timestamp: {result.timestamp.strftime('%Y-%m-%d %H:%M')}")

            # Inputs summary
            summary_parts.append("   Inputs:")
            for name, value in result.inputs.items():
                formatted = self._format_value(value)
                summary_parts.append(f"      - {name}: {formatted}")

            # Outputs summary
            summary_parts.append("   Outputs:")
            for name, value in result.outputs.items():
                formatted = self._format_value(value)
                summary_parts.append(f"      - {name}: {formatted}")

            summary_parts.append("")

        summary_parts.append("=" * 60)
        summary_parts.append("END OF SUMMARY")
        summary_parts.append("=" * 60)

        return "\n".join(summary_parts)


# Module exports
__all__ = [
    "ReportService",
    "ReportOptions",
]
